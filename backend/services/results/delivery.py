from __future__ import annotations

import hashlib
import re
import smtplib
import ssl
from email.message import EmailMessage
from email.utils import getaddresses
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches
from flask import current_app

from ...core.database import Database, utcnow
from ..security import SecretVault
from .manifests import ResultService


ARTIFACT_KINDS = ("summary_docx", "report_docx", "dashboard_png")


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _published(database: Database, run_id: str, workspace_id: str) -> tuple[dict, dict]:
    service = ResultService(database)
    publication = service.publication(run_id, workspace_id=workspace_id)
    if not publication:
        raise PermissionError("未通过验证发布门禁，不得导出或发送正式成果")
    manifest = service.manifest(publication["manifest_id"], workspace_id=workspace_id)
    if not manifest or manifest.get("status") != "published":
        raise PermissionError("发布版本已失效")
    return publication, manifest


def _record_artifact(
    database: Database, path: Path, kind: str, workspace_id: str,
    run_id: str, publication: dict, manifest: dict,
) -> dict[str, Any]:
    digest = _sha256(path)
    return database.put("artifacts", {
        "id": database.new_id("art"), "workspace_id": workspace_id,
        "run_id": run_id, "publication_id": publication["id"], "manifest_id": manifest["id"],
        "manifest_version": manifest["version"], "title": path.stem, "kind": kind,
        "filename": path.name, "path": str(path), "size": path.stat().st_size,
        "sha256": digest, "status": "ready", "immutable": True,
    }, workspace_id=workspace_id)


def _doc_header(document: Document, manifest: dict, title: str) -> dict[str, Any]:
    payload = manifest["payload"]
    contract = payload["contract"]
    document.add_heading(title, 0)
    document.add_paragraph(f"成果版本：{manifest['version']}  生成时间：{manifest['created_at']}")
    document.add_heading("业务分析目标", level=1)
    document.add_paragraph(str(contract.get("objective") or ""))
    document.add_heading("统计覆盖范围", level=1)
    document.add_paragraph(str(contract.get("coverage") or ""))
    return payload


def _write_kpis(document: Document, kpis: list[dict[str, Any]]) -> None:
    document.add_heading("四项关键指标", level=1)
    table = document.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    for index, item in enumerate(kpis[:4]):
        table.cell(0, index).text = str(item.get("label") or "不可用")
        value = item.get("value")
        table.cell(1, index).text = str(value) if value is not None else str(item.get("unavailable_reason") or "不可用")


def _render_png(path: Path, charts: list[dict[str, Any]]) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    figure, axes = plt.subplots(2, 2, figsize=(14, 9), constrained_layout=True)
    for index, (axis, chart) in enumerate(zip(axes.flat, charts[:4])):
        option = chart.get("option") or {}
        title = ["Value and trend", "Composition", "Stacked comparison", "Category comparison"][index]
        axis.set_title(title)
        if not chart.get("available"):
            axis.text(0.5, 0.5, "No applicable verified data", ha="center", va="center")
            axis.set_axis_off()
            continue
        series = option.get("series") or []
        labels = (option.get("xAxis") or {}).get("data") or []
        if chart.get("type") == "pie":
            data = series[0].get("data") or [] if series else []
            values = [float(item["value"]) for item in data if item.get("value") is not None]
            names = [str(item["name"])[:20] for item in data if item.get("value") is not None]
            if values and sum(abs(value) for value in values) > 0 and all(value >= 0 for value in values):
                axis.pie(values, labels=names, autopct="%1.1f%%")
            else:
                axis.text(0.5, 0.5, "Pie requires non-negative values", ha="center", va="center")
            continue
        positions = list(range(len(labels)))
        if chart.get("type") == "bar_line" and series:
            axis.bar(positions, series[0].get("data") or [], color="#2563eb")
            twin = axis.twinx()
            twin.plot(positions, series[1].get("data") or [], color="#f97316", marker="o")
        elif chart.get("type") == "stacked_bar" and series:
            bottom = [0.0] * len(positions)
            for color, current in zip(("#2563eb", "#14b8a6"), series):
                values = [float(value or 0) for value in current.get("data") or []]
                axis.bar(positions, values, bottom=bottom, color=color, label=current.get("name"))
                bottom = [left + right for left, right in zip(bottom, values)]
            axis.legend()
        elif series:
            axis.bar(positions, series[0].get("data") or [], color="#2563eb")
        axis.set_xticks(positions, [str(value)[:16] for value in labels], rotation=30, ha="right")
        axis.grid(axis="y", alpha=0.2)
    figure.suptitle("Verified analysis dashboard", fontsize=16)
    figure.savefig(path, dpi=160, facecolor="white")
    plt.close(figure)


def generate_artifact(database: Database, run_id: str, workspace_id: str, kind: str) -> dict[str, Any]:
    if kind not in ARTIFACT_KINDS:
        raise ValueError("成果类型必须是 summary_docx、report_docx 或 dashboard_png")
    publication, manifest = _published(database, run_id, workspace_id)
    existing = next((
        item for item in database.list("artifacts", workspace_id=workspace_id, limit=5000)
        if item.get("manifest_id") == manifest["id"] and item.get("kind") == kind and item.get("status") == "ready"
        and Path(str(item.get("path") or "")).is_file()
    ), None)
    if existing:
        return existing
    export_dir = current_app.config["SETTINGS"].export_dir
    base = f"{manifest['id']}_{manifest['version']}"
    payload = manifest["payload"]
    if kind == "dashboard_png":
        path = export_dir / f"{base}_dashboard.png"
        _render_png(path, payload.get("charts") or [])
    else:
        path = export_dir / f"{base}_{'summary' if kind == 'summary_docx' else 'report'}.docx"
        document = Document()
        current = _doc_header(
            document, manifest,
            "极简分析结论" if kind == "summary_docx" else "完整数据分析报告",
        )
        document.add_heading("经验证结论", level=1)
        document.add_paragraph(str(current.get("summary") or ""))
        _write_kpis(document, current.get("kpis") or [])
        if kind == "report_docx":
            png_path = export_dir / f"{base}_dashboard_embed.png"
            _render_png(png_path, current.get("charts") or [])
            document.add_heading("四图看板", level=1)
            document.add_picture(str(png_path), width=Inches(6.4))
            report = current.get("report") or {}
            document.add_heading("数据结果与归因边界", level=1)
            for item in report.get("attribution") or []:
                document.add_paragraph(f"[{item.get('type', 'fact')}] {item.get('text', '')}")
            document.add_heading("建议", level=1)
            recommendations = report.get("recommendations") or {}
            for key, label in (("short_term", "短期"), ("medium_term", "中期"), ("long_term", "长期")):
                values = recommendations.get(key) or []
                document.add_paragraph(f"{label}：" + ("；".join(str(value) for value in values) if values else "尚无经证据支持的建议"))
            document.add_heading("限制与存疑问题", level=1)
            for value in current.get("limitations") or ["无"]:
                document.add_paragraph(str(value), style="List Bullet")
        document.save(path)
    return _record_artifact(database, path, kind, workspace_id, run_id, publication, manifest)


def generate_artifacts(database: Database, run_id: str, workspace_id: str, kinds: list[str] | None = None) -> list[dict]:
    requested = list(dict.fromkeys(kinds or ARTIFACT_KINDS))
    return [generate_artifact(database, run_id, workspace_id, kind) for kind in requested]


def _recipients(value: str | list[str]) -> list[str]:
    raw = ",".join(value) if isinstance(value, list) else str(value or "")
    addresses = [address for _name, address in getaddresses([raw.replace("，", ",")]) if address]
    pattern = re.compile(r"^[^@\s,]+@[^@\s,]+\.[^@\s,]+$")
    if not addresses or len(addresses) > 50 or any(not pattern.fullmatch(item) for item in addresses):
        raise ValueError("收件人必须是有效的逗号分隔邮箱，最多 50 个")
    return list(dict.fromkeys(addresses))


def _email_message(
    *, sender: str, recipients: list[str], subject: str, text: str, artifacts: list[dict],
) -> EmailMessage:
    message = EmailMessage()
    message["Subject"] = subject[:200]
    message["From"] = sender
    message["To"] = ", ".join(recipients)
    message.set_content(text)
    for artifact in artifacts:
        path = Path(artifact["path"])
        subtype = "vnd.openxmlformats-officedocument.wordprocessingml.document" if path.suffix == ".docx" else "png"
        maintype = "application" if path.suffix == ".docx" else "image"
        message.add_attachment(path.read_bytes(), maintype=maintype, subtype=subtype, filename=artifact["filename"])
    return message


def prepare_eml(
    database: Database, run_id: str, workspace_id: str, *, recipients: str | list[str],
    subject: str, text: str, kinds: list[str] | None = None, sender: str = "analysis@localhost",
) -> tuple[dict, list[dict]]:
    publication, manifest = _published(database, run_id, workspace_id)
    addresses = _recipients(recipients)
    artifacts = generate_artifacts(database, run_id, workspace_id, kinds)
    message = _email_message(
        sender=sender, recipients=addresses, subject=subject, text=text, artifacts=artifacts,
    )
    signature = hashlib.sha256((
        manifest["id"] + "\0" + ",".join(addresses) + "\0" + subject + "\0" +
        "\0".join(item["sha256"] for item in artifacts)
    ).encode("utf-8")).hexdigest()[:20]
    path = current_app.config["SETTINGS"].export_dir / f"{manifest['id']}_{signature}.eml"
    path.write_bytes(message.as_bytes())
    eml = _record_artifact(database, path, "email_eml", workspace_id, run_id, publication, manifest)
    eml["recipients"] = addresses
    return eml, artifacts


def send_email(
    database: Database, run_id: str, workspace_id: str, *, connector: dict,
    recipients: str | list[str], subject: str, text: str, kinds: list[str] | None,
    idempotency_key: str,
) -> dict[str, Any]:
    if connector.get("type") != "email" or not connector.get("enabled", True):
        raise ValueError("请选择已启用的 SMTP 邮件连接器")
    if not idempotency_key:
        raise ValueError("邮件发送必须提供 Idempotency-Key")
    existing = next((
        item for item in database.list("email_deliveries", workspace_id=workspace_id, limit=5000)
        if item.get("idempotency_key") == idempotency_key and item.get("run_id") == run_id
    ), None)
    if existing:
        return existing
    secret = SecretVault(current_app.config["VAULT_KEY"]).open(connector.get("credential", ""), {})
    sender = str(secret.get("sender") or secret.get("username") or "")
    if not sender or not secret.get("host"):
        raise ValueError("SMTP 连接器缺少 host/sender")
    eml, artifacts = prepare_eml(
        database, run_id, workspace_id, recipients=recipients, subject=subject, text=text,
        kinds=kinds, sender=sender,
    )
    addresses = eml["recipients"]
    delivery = database.put("email_deliveries", {
        "id": database.new_id("mail"), "workspace_id": workspace_id, "run_id": run_id,
        "publication_id": eml["publication_id"], "manifest_id": eml["manifest_id"],
        "connector_id": connector["id"], "idempotency_key": idempotency_key,
        "status": "sending", "recipients": addresses, "subject": subject,
        "artifact_ids": [item["id"] for item in artifacts],
        "artifact_hashes": {item["id"]: item["sha256"] for item in artifacts},
        "eml_artifact_id": eml["id"], "started_at": utcnow(),
    }, workspace_id=workspace_id)
    message = _email_message(sender=sender, recipients=addresses, subject=subject, text=text, artifacts=artifacts)
    try:
        if secret.get("use_ssl"):
            client = smtplib.SMTP_SSL(secret["host"], int(secret.get("port", 465)), timeout=30, context=ssl.create_default_context())
        else:
            client = smtplib.SMTP(secret["host"], int(secret.get("port", 587)), timeout=30)
        with client as smtp:
            if secret.get("use_tls", True) and not secret.get("use_ssl"):
                smtp.starttls(context=ssl.create_default_context())
            if secret.get("username"):
                smtp.login(secret["username"], secret.get("password", ""))
            refused = smtp.send_message(message, to_addrs=addresses)
        status = "sent" if not refused else "partial"
        return database.patch("email_deliveries", delivery["id"], {
            "status": status, "refused": refused, "finished_at": utcnow(),
        }, workspace_id=workspace_id) or delivery
    except Exception as exc:
        # Network failure after DATA can have an unknown delivery outcome. Idempotency
        # prevents automatic resend; a user may explicitly choose a new key.
        database.patch("email_deliveries", delivery["id"], {
            "status": "unknown", "error": str(exc), "finished_at": utcnow(),
        }, workspace_id=workspace_id)
        raise ConnectionError("邮件发送状态未知，系统不会使用同一幂等键自动重发") from exc
