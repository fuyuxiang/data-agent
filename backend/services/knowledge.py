from __future__ import annotations

import hashlib
import json
import re
import shutil
from collections import Counter
from pathlib import Path
import pandas as pd
from docx import Document
from flask import current_app
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename

from ..core.database import Database


TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".html", ".docx", ".xlsx", ".xls", ".pdf"}


def _db() -> Database:
    return current_app.extensions["meridian_db"]


def _extract(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".html", ".json"}:
        text = path.read_text(encoding="utf-8", errors="replace")
        return "\n".join(f"[line {index}] {line}" for index, line in enumerate(text.splitlines(), 1))
    if suffix == ".csv":
        try:
            frame = pd.read_csv(path)
        except UnicodeDecodeError:
            frame = pd.read_csv(path, encoding="gb18030")
        return frame.to_csv(index=False)
    if suffix == ".docx":
        document = Document(path)
        parts = [
            f"[paragraph {index}] {paragraph.text}"
            for index, paragraph in enumerate(document.paragraphs, 1) if paragraph.text.strip()
        ]
        for table_index, table in enumerate(document.tables, 1):
            for row_index, row in enumerate(table.rows, 1):
                for column_index, cell in enumerate(row.cells, 1):
                    value = cell.text.strip()
                    if value:
                        parts.append(f"[table {table_index} cell R{row_index}C{column_index}] {value}")
        return "\n".join(parts)
    if suffix == ".pdf":
        from pypdf import PdfReader

        parts = []
        for index, page in enumerate(PdfReader(path).pages, 1):
            text = (page.extract_text() or "").strip()
            parts.append(
                f"[page {index}] {text}" if text
                else f"[page {index}] NO_EXTRACTABLE_TEXT; VISUAL_EVIDENCE_REQUIRED"
            )
        return "\n\n".join(parts)
    if suffix in {".xlsx", ".xls"}:
        sheets = pd.read_excel(path, sheet_name=None)
        parts = []
        for name, frame in sheets.items():
            parts.append(f"[sheet {name} header] " + " | ".join(str(value) for value in frame.columns))
            for row_index, row in frame.head(current_app.config["SETTINGS"].max_ingest_rows).iterrows():
                parts.append(
                    f"[sheet {name} row {int(row_index) + 2}] "
                    + " | ".join(f"{column}={row[column]}" for column in frame.columns)
                )
        return "\n".join(parts)
    raise ValueError("该文档格式暂不支持文本解析")


def _chunks(text: str, size: int = 1200, overlap: int = 160) -> list[str]:
    normalized = re.sub(r"\r\n?", "\n", text)
    paragraphs = [item.strip() for item in re.split(r"\n{2,}", normalized) if item.strip()]
    chunks: list[str] = []
    buffer = ""
    for paragraph in paragraphs:
        if len(buffer) + len(paragraph) + 2 <= size:
            buffer = f"{buffer}\n\n{paragraph}".strip()
            continue
        if buffer:
            chunks.append(buffer)
        tail = buffer[-overlap:] if buffer else ""
        buffer = f"{tail}\n{paragraph}".strip()
        while len(buffer) > size:
            chunks.append(buffer[:size])
            buffer = buffer[size - overlap:]
    if buffer:
        chunks.append(buffer)
    return chunks or ([normalized[:size]] if normalized else [])


def _tokens(text: str) -> list[str]:
    lowered = str(text or "").lower()
    words = re.findall(r"[a-z0-9_]+", lowered)
    cjk_runs = re.findall(r"[\u4e00-\u9fff]+", lowered)
    grams: list[str] = []
    for run in cjk_runs:
        grams.extend(list(run))
        for size in (2, 3, 4):
            grams.extend(run[index:index + size] for index in range(max(0, len(run) - size + 1)))
    return words + grams


def _embedding(text: str, workspace_id: str = "default") -> list[float]:
    from .embeddings import embed

    return embed(text, workspace_id)


def _cosine(left: list[float], right: list[float]) -> float:
    if not left or len(left) != len(right):
        # Indexes created by a previous provider must be rebuilt; never score
        # truncated vectors from different embedding spaces as comparable.
        return 0.0
    return sum(a * b for a, b in zip(left, right))


def _lexical_score(query_tokens: Counter, content_tokens: list[str], query: str, content: str) -> float:
    document = Counter(content_tokens)
    overlap = sum(min(count, document.get(token, 0)) for token, count in query_tokens.items())
    coverage = overlap / max(1, sum(query_tokens.values()))
    phrase = 1.0 if query.lower().strip() in content.lower() else 0.0
    return coverage + phrase


def _build_chunk_index(chunks: list[str], workspace_id: str = "default") -> list[dict]:
    return [
        {
            "index": index, "content_hash": hashlib.sha256(chunk.encode("utf-8")).hexdigest(),
            "tokens": _tokens(chunk), "embedding": _embedding(chunk, workspace_id),
        }
        for index, chunk in enumerate(chunks)
    ]


def add_document(file: FileStorage, workspace_id: str, tags: list[str] | None = None) -> dict:
    filename = secure_filename(file.filename or "")
    suffix = Path(filename).suffix.lower()
    if not filename or suffix not in TEXT_EXTENSIONS:
        raise ValueError("支持 TXT、Markdown、HTML、CSV、JSON、PDF、Word 和 Excel 文档")
    document_id = _db().new_id("doc")
    target = current_app.config["SETTINGS"].knowledge_dir / f"{document_id}{suffix}"
    file.save(target)
    try:
        if target.stat().st_size > 50 * 1024 * 1024:
            raise ValueError("分析入口单文件不得超过 50MB")
        text = _extract(target)
    except Exception:
        target.unlink(missing_ok=True)
        raise
    if not text.strip():
        target.unlink(missing_ok=True)
        raise ValueError("文档中没有可检索文本")
    chunks = _chunks(text)
    record = _db().put(
        "knowledge_documents",
        {
            "id": document_id, "workspace_id": workspace_id, "name": Path(filename).stem,
            "filename": filename, "format": suffix.lstrip("."), "path": str(target), "text": text,
            "chunks": chunks, "chunk_index": _build_chunk_index(chunks, workspace_id), "tags": tags or [],
            "enabled": True, "characters": len(text), "chunk_count": len(chunks), "index_version": 2,
            "evidence_locations": True,
            "visual_only_pages": [
                int(value) for value in re.findall(r"\[page (\d+)\] NO_EXTRACTABLE_TEXT", text)
            ],
        },
        workspace_id=workspace_id,
    )
    return public_document(record)


def index_document_path(
    path: Path, workspace_id: str, *, original_name: str = "", tags: list[str] | None = None,
    import_filename: str = "",
) -> dict:
    """Commit an already uploaded preview file to the searchable index."""
    suffix = path.suffix.lower()
    if suffix not in TEXT_EXTENSIONS or not path.is_file():
        raise ValueError("待确认的知识文件不存在或格式不支持")
    text = _extract(path)
    if not text.strip():
        raise ValueError("文档中没有可检索文本")
    document_id = _db().new_id("doc")
    target = current_app.config["SETTINGS"].knowledge_dir / f"{document_id}{suffix}"
    if path.resolve() != target.resolve():
        shutil.copyfile(path, target)
    chunks = _chunks(text)
    filename = original_name or path.name
    record = _db().put(
        "knowledge_documents",
        {
            "id": document_id, "workspace_id": workspace_id, "name": Path(filename).stem,
            "filename": filename, "format": suffix.lstrip("."), "path": str(target), "text": text,
            "chunks": chunks, "chunk_index": _build_chunk_index(chunks, workspace_id), "tags": tags or [],
            "enabled": True, "characters": len(text), "chunk_count": len(chunks), "index_version": 2,
            "source_import_filename": import_filename,
        },
        workspace_id=workspace_id,
    )
    return public_document(record)


METRIC_HEADERS = {
    "name": {"name", "名称", "指标名", "指标名称", "指标", "metric"},
    "alias": {"alias", "别名", "别称", "aliases"},
    "definition": {"definition", "定义", "desc", "description", "说明"},
    "sql_template": {"sql", "sql_template", "sql模板", "query"},
    "notes": {"notes", "备注", "note", "remark", "说明2"},
}
RULE_HEADERS = {
    "rule_id": {"rule_id", "规则id", "id", "rule"},
    "description": {"description", "描述", "说明", "desc"},
    "condition": {"condition", "条件", "断言", "assert"},
    "severity": {"severity", "严重程度", "level", "等级"},
}
NOTE_HEADERS = {
    "topic": {"topic", "主题", "话题", "subject"},
    "content": {"content", "内容", "text", "正文"},
    "tags": {"tags", "标签", "tag", "关键词"},
}
KNOWLEDGE_TEMPLATES = {
    "metrics": METRIC_HEADERS, "business_rules": RULE_HEADERS, "context_notes": NOTE_HEADERS,
}


def _structured_frame(frame: pd.DataFrame) -> tuple[str, list[dict]] | None:
    normalized = {str(column).strip().lower(): str(column) for column in frame.columns}
    best_type, best_headers, best_score = "", {}, 0
    for table_type, headers in KNOWLEDGE_TEMPLATES.items():
        mapped = {
            field: next((normalized[alias] for alias in aliases if alias in normalized), None)
            for field, aliases in headers.items()
        }
        score = sum(value is not None for value in mapped.values())
        if score > best_score:
            best_type, best_headers, best_score = table_type, mapped, score
    if best_score < 2:
        return None
    records = []
    for _, row in frame.iterrows():
        record = {"table": best_type}
        for field, column in best_headers.items():
            value = row.get(column) if column else None
            record[field] = str(value).strip() if column and pd.notna(value) else ""
        if any(value for key, value in record.items() if key != "table"):
            records.append(record)
    return best_type, records


def _records_from_llm(text: str, provider_id: str, workspace_id: str) -> list[dict]:
    from .models import resolve_provider
    from .usage import ensure_quota, record_usage, response_usage

    provider, client = resolve_provider(provider_id or None, workspace_id)
    if not provider or not client:
        raise ValueError("未配置可用的 LLM 模型，无法解析非结构化知识")
    prompt = """从以下文本提取业务指标、业务规则和背景知识。只输出 JSON：
{"metrics":[{"name":"","alias":"","definition":"","sql_template":"","notes":""}],"business_rules":[{"rule_id":"","description":"","condition":"","severity":"warning"}],"context_notes":[{"topic":"","content":"","tags":""}]}
无法提取的字段留空，不得编造。
文本：
""" + text[:48_000]
    quota = ensure_quota(_db(), workspace_id)
    max_tokens = min(
        int(provider.get("max_output_tokens") or current_app.config["SETTINGS"].default_max_output_tokens),
        quota["remaining"], 8192,
    )
    response = client.chat.completions.create(
        model=provider["model"], messages=[{"role": "user", "content": prompt}],
        temperature=0, max_tokens=max(1, max_tokens),
    )
    record_usage(_db(), workspace_id, response_usage(response, provider["model"]), operation="knowledge_extraction")
    raw = str(response.choices[0].message.content or "")
    raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw.strip(), flags=re.IGNORECASE)
    try:
        payload = json.loads(raw)
    except json.JSONDecodeError:
        match = re.search(r"\{[\s\S]*\}", raw)
        payload = json.loads(match.group()) if match else {}
    if not isinstance(payload, dict):
        return []
    result = []
    for table_type in KNOWLEDGE_TEMPLATES:
        for item in payload.get(table_type, []):
            if isinstance(item, dict):
                result.append({"table": table_type, **item})
    return result


def parse_knowledge_path(path: Path, workspace_id: str, provider_id: str = "") -> dict:
    if path.suffix.lower() in {".xlsx", ".xls"}:
        sheets = pd.read_excel(path, sheet_name=None)
        structured, unstructured = [], []
        for name, frame in sheets.items():
            if frame.empty:
                continue
            parsed = _structured_frame(frame)
            if parsed:
                structured.extend(parsed[1])
            else:
                unstructured.append(f"[Sheet: {name}]\n{frame.to_csv(index=False)}")
        preview = list(structured)
        if unstructured:
            preview.extend(_records_from_llm("\n\n".join(unstructured), provider_id, workspace_id))
        format_name = "mixed" if structured and unstructured else "structured" if structured else "unstructured"
        return {"format": format_name, "preview": preview}
    if path.suffix.lower() == ".docx":
        return {"format": "unstructured", "preview": _records_from_llm(_extract(path), provider_id, workspace_id)}
    raise ValueError("只支持 .xlsx、.xls 和 .docx 知识导入")


def strip_temp_prompt_thinking(value: object) -> str:
    text = str(value or "").strip()
    text = re.sub(r"<think\b[^>]*>[\s\S]*?</think\s*>", "", text, flags=re.IGNORECASE)
    unclosed = re.search(r"<think\b[^>]*>", text, flags=re.IGNORECASE)
    if unclosed:
        text = text[:unclosed.start()]
    text = re.sub(r"</think\s*>", "", text, flags=re.IGNORECASE)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def public_document(document: dict) -> dict:
    value = dict(document)
    for key in ("text", "chunks", "chunk_index", "path"):
        value.pop(key, None)
    return value


def _entry_text(entry: dict) -> str:
    entry_type = entry.get("type")
    if entry_type == "metric":
        fields = ("name", "alias", "definition", "sql_template", "notes")
    elif entry_type == "business_rule":
        fields = ("name", "description", "condition", "severity")
    else:
        fields = ("name", "topic", "content", "tags")
    values = []
    for field in fields:
        value = entry.get(field)
        if isinstance(value, list):
            value = " ".join(str(item) for item in value)
        if value:
            values.append(str(value))
    return "\n".join(values)


def _search_rows(workspace_id: str) -> list[dict]:
    rows = []
    for document in _db().list("knowledge_documents", workspace_id=workspace_id):
        if not document.get("enabled", True):
            continue
        chunks = document.get("chunks") or []
        index = document.get("chunk_index")
        if not isinstance(index, list) or len(index) != len(chunks):
            index = _build_chunk_index(chunks, workspace_id)
            document.update({"chunk_index": index, "index_version": 2})
            _db().put("knowledge_documents", document, workspace_id=workspace_id)
        for chunk_number, (chunk, indexed) in enumerate(zip(chunks, index)):
            rows.append({
                "id": f"{document['id']}:{chunk_number}", "document_id": document["id"],
                "document_name": document["name"], "chunk": chunk_number, "text": chunk,
                "tokens": indexed.get("tokens") or _tokens(chunk),
                "embedding": indexed.get("embedding") or _embedding(chunk, workspace_id), "kind": "document",
            })
    for entry in _db().list("knowledge_entries", workspace_id=workspace_id):
        if not entry.get("enabled", True):
            continue
        text = _entry_text(entry)
        rows.append({
            "id": entry["id"], "document_id": entry["id"],
            "document_name": entry.get("name") or entry.get("topic") or entry["id"],
            "chunk": 0, "text": text, "tokens": entry.get("tokens") or _tokens(text),
            "embedding": entry.get("embedding") or _embedding(text, workspace_id), "kind": entry.get("type"),
        })
    return rows


def _ranks(rows: list[dict], key: str) -> dict[str, int]:
    ordered = sorted(rows, key=lambda item: float(item.get(key, 0)), reverse=True)
    return {item["id"]: rank for rank, item in enumerate(ordered, 1)}


def search(
    query: str, workspace_id: str, limit: int = 6,
    document_ids: list[str] | tuple[str, ...] | set[str] | None = None,
) -> list[dict]:
    rows = _search_rows(workspace_id)
    if document_ids is not None:
        allowed = {str(value) for value in document_ids}
        rows = [row for row in rows if row.get("document_id") in allowed]
    if not rows or not query.strip():
        return []
    query_embedding = _embedding(query, workspace_id)
    query_tokens = Counter(_tokens(query))
    for row in rows:
        row["vector_score"] = max(0.0, _cosine(query_embedding, row["embedding"]))
        row["lexical_score"] = _lexical_score(query_tokens, row["tokens"], query, row["text"])
    vector_ranks = _ranks(rows, "vector_score")
    lexical_ranks = _ranks(rows, "lexical_score")
    for row in rows:
        row["rrf_score"] = 1 / (60 + vector_ranks[row["id"]]) + 1 / (60 + lexical_ranks[row["id"]])
        row["score"] = row["rrf_score"] + row["vector_score"] + min(row["lexical_score"], 2.0) * 0.5
    relevant = [item for item in rows if item["vector_score"] >= 0.08 or item["lexical_score"] > 0]
    relevant.sort(key=lambda item: item["score"], reverse=True)
    return [
        {
            key: (round(float(value), 6) if key.endswith("score") else value)
            for key, value in item.items()
            if key not in {"tokens", "embedding", "id"}
        }
        for item in relevant[: max(1, min(limit, 20))]
    ]


def save_entry(payload: dict, workspace_id: str, entry_id: str | None = None) -> dict:
    current = _db().get("knowledge_entries", entry_id) if entry_id else None
    entry_type = str(payload.get("type") or (current or {}).get("type") or "context_note")
    if entry_type not in {"metric", "business_rule", "context_note"}:
        raise ValueError("知识条目类型必须是 metric、business_rule 或 context_note")
    name = str(payload.get("name") or payload.get("topic") or "").strip()
    if not name:
        raise ValueError("知识条目名称不能为空")
    if current and current.get("workspace_id", "default") != workspace_id:
        raise PermissionError("知识条目不属于当前工作空间")
    record = {
        **(current or {}), **payload,
        "id": entry_id or _db().new_id("kb"), "workspace_id": workspace_id,
        "type": entry_type, "name": name[:160], "enabled": bool(payload.get("enabled", True)),
    }
    text = _entry_text(record)
    record.update({
        "content_hash": hashlib.sha256(text.encode("utf-8")).hexdigest(),
        "tokens": _tokens(text), "embedding": _embedding(text, workspace_id),
    })
    return _db().put("knowledge_entries", record, workspace_id=workspace_id)
